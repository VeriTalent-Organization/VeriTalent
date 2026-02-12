"""
Document Loader

Utilities for loading and extracting text from various document formats.
"""
import io
from typing import Optional


async def load_document(
    content: bytes,
    filename: str,
    content_type: str,
) -> str:
    """
    Load a document and extract text content.
    
    Args:
        content: Raw file content
        filename: Original filename
        content_type: MIME type
        
    Returns:
        Extracted text content
    """
    if content_type == "application/pdf":
        return await extract_pdf_text(content)
    elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return await extract_docx_text(content)
    elif content_type == "text/plain":
        return content.decode("utf-8", errors="ignore")
    else:
        # Try to decode as text
        return content.decode("utf-8", errors="ignore")


async def extract_pdf_text(content: bytes) -> str:
    """
    Extract text from PDF content.
    
    Args:
        content: PDF file content
        
    Returns:
        Extracted text
    """
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(stream=content, filetype="pdf")
        text_parts = []
        
        for page in doc:
            text_parts.append(page.get_text())
        
        doc.close()
        return "\n\n".join(text_parts)
        
    except ImportError:
        # Fallback to pdfplumber
        try:
            import pdfplumber
            
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                
                return "\n\n".join(text_parts)
                
        except ImportError:
            raise RuntimeError("No PDF library available. Install pymupdf or pdfplumber.")


async def extract_docx_text(content: bytes) -> str:
    """
    Extract text from DOCX content.
    
    Args:
        content: DOCX file content
        
    Returns:
        Extracted text
    """
    try:
        from docx import Document
        
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return "\n\n".join(text_parts)
        
    except ImportError:
        raise RuntimeError("python-docx not installed. Install with: pip install python-docx")


def detect_file_type(content: bytes, filename: str) -> str:
    """
    Detect file type from content and filename.
    
    Args:
        content: File content
        filename: Filename
        
    Returns:
        MIME type
    """
    # Check magic bytes
    if content[:4] == b"%PDF":
        return "application/pdf"
    elif content[:4] == b"PK\x03\x04":  # ZIP-based (DOCX, XLSX, etc.)
        if filename.lower().endswith(".docx"):
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif filename.lower().endswith(".xlsx"):
            return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    
    # Fallback to extension
    ext = filename.lower().split(".")[-1] if "." in filename else ""
    
    ext_map = {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "doc": "application/msword",
        "txt": "text/plain",
        "md": "text/markdown",
    }
    
    return ext_map.get(ext, "application/octet-stream")
